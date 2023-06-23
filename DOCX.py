from docx import Document
import langid
from bs4 import BeautifulSoup
import re

# Get title
def title_docx(document):
    doc = Document(document)
    title = doc.core_properties.title
    return title

# Count pages
def page_count_docx(document):
    doc = Document(document)
    return len(doc.sections)

# Count words
def word_count_docx(document):
    doc = Document(document)
    words = 0
    for paragraph in doc.paragraphs:
        words += len(paragraph.text.split())
    return words

# Extract keywords
def keywords_docx(document):
    doc = Document(document)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # Remove non-alphanumeric characters and convert to lowercase
    cleaned_text = re.sub(r'[^a-zA-Z0-9]+', ' ', text.lower())

    # Split text into individual words
    words = cleaned_text.split()

    # Count word frequency, excluding numbers
    word_count = {}
    for word in words:
        if word.isalpha():  # Exclude words that consist only of alphabetic characters
            if word in word_count:
                word_count[word] += 1
            else:
                word_count[word] = 1

    # Sort words by frequency in descending order
    sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)

    # Extract the top 3 most frequent words as keywords
    keywords = [word for word, _ in sorted_words[:3]]

    return keywords

# Count characters 
def character_count_docx(document):
    doc = Document(document)
    num_characters = 0
    for paragraph in doc.paragraphs:
        num_characters += len(paragraph.text)
    return num_characters


# Get typography
def typography_docx(document):
    doc = Document(document)
    typography = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font_name = run.font.name
            if font_name is None:
                font_name = "Unknown"
            typography.add(font_name)
    return list(typography)

# Count images
def image_count_docx(document):
    doc = Document(document)
    images = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images += 1
    return images

# Count tables
def table_count_docx(document):
    with open(document, 'rb') as f:
        soup = BeautifulSoup(f, 'xml')
        tables = soup.find_all('w:tbl')
        return len(tables)

# Detect language
def language_docx(document):
    doc = Document(document)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text
    language = langid.classify(text)[0]
    return language

# Detect encryption
def encryption_docx(document):
    doc = Document(document)
    settings = doc.settings
    if hasattr(settings, 'password') and settings.password:
        return True
    else:
        return False
    
# Get author
def author_docx(document):
    doc = Document(document)
    core_props = doc.core_properties
    if core_props.author:
        return core_props.author
    return ""